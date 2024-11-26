import pandas as pd
import numpy as np
import statsmodels.api as sm
from statsmodels.iolib.summary2 import summary_col  # Ensure this import is present

# Load the data
file_path = 'filtered_with_FTA_columns_asean.xlsx'
data = pd.read_excel(file_path)

# Drop all rows with NaNs
data = data.dropna()

# Take the natural log of relevant variables
data['ln_X_ijt'] = np.log(data['tradeflow_comtrade_o'])  # Assuming 'tradeflow_comtrade_o' is X_ijt
data['ln_Y_it'] = np.log(data['gdp_o'])  # GDP of origin country i
data['ln_Y_jt'] = np.log(data['gdp_d'])  # GDP of destination country j
data['ln_Pop_it'] = np.log(data['pop_o'])  # Population of origin country i
data['ln_Pop_jt'] = np.log(data['pop_d'])  # Population of destination country j
data['ln_Dist_ij'] = np.log(data['distw_harmonic'])  # Assuming 'distw_harmonic' is distance between countries i and j

# Base independent variables
base_vars = ['ln_Y_it', 'ln_Y_jt', 'ln_Pop_it', 'ln_Pop_jt', 'ln_Dist_ij', 'comlang_off', 'contig']

# Regressions
X1 = data[base_vars + ['FTA1_ijt']]
X1 = sm.add_constant(X1)  # Add a constant term
model1 = sm.OLS(data['ln_X_ijt'], X1).fit(cov_type='HC3')

X2 = data[base_vars + ['FTA1_ijt', 'FTA2_ict']]
X2 = sm.add_constant(X2)
model2 = sm.OLS(data['ln_X_ijt'], X2).fit(cov_type='HC3')

X3 = data[base_vars + ['FTA1_ijt', 'FTA2_ict', 'FTA3_cjt']]
X3 = sm.add_constant(X3)
model3 = sm.OLS(data['ln_X_ijt'], X3).fit(cov_type='HC3')

# Combine results into one table
results_table = summary_col([model1, model2, model3], 
                            stars=True, 
                            model_names=["FTA1 Only", "FTA1 & FTA2", "FTA1, FTA2 & FTA3"],
                            info_dict={'R-squared': lambda x: f"{x.rsquared:.4f}",
                                       'No. observations': lambda x: f"{int(x.nobs)}"})

# Save results in LaTeX
latex_path = '../output/regression/regression_results_fta_asean.tex'
with open(latex_path, 'w') as f:
    f.write(results_table.as_latex())
    
# Save results in DOCX
import docx
doc = docx.Document()
doc.add_paragraph("Regression Results for FTA Models")
doc.add_paragraph(results_table.as_text())
docx_path = '../output/regression/regression_results_fta_asean.docx'
doc.save(docx_path)

print("Regression results saved in LaTeX, PDF, and DOCX formats.")