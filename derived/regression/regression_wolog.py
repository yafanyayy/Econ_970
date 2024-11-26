import pandas as pd
import numpy as np
import statsmodels.api as sm
from statsmodels.iolib.summary2 import summary_col  # Ensure this import is present

# Load the data
file_path = 'filtered_with_FTA_columns_asean.xlsx'
data = pd.read_excel(file_path)

# Drop all rows with NaNs
data = data.dropna()

# Define the dependent variable
Y = data['tradeflow_comtrade_o']  # Assuming 'tradeflow_comtrade_o' is X_ijt

# # Ensure 'country_id_o', 'country_id_d', and 'year' are categorical variables
# data['country_id_o'] = data['country_id_o'].astype('category')
# data['country_id_d'] = data['country_id_d'].astype('category')
# data['year'] = data['year'].astype('category')

# # First regression: FTA1 only
# formula1 = 'tradeflow_comtrade_o ~ FTA1_ijt + C(country_id_o) + C(country_id_d) + C(year)'
# model1 = sm.OLS.from_formula(formula1, data=data).fit(cov_type='HC3')

# # Second regression: FTA1 and FTA2
# formula2 = 'tradeflow_comtrade_o ~ FTA1_ijt + FTA2_ict + C(country_id_o) + C(country_id_d) + C(year)'
# model2 = sm.OLS.from_formula(formula2, data=data).fit(cov_type='HC3')

# # Third regression: FTA1, FTA2, and FTA3
# formula3 = 'tradeflow_comtrade_o ~ FTA1_ijt + FTA2_ict + FTA3_cjt + C(country_id_o) + C(country_id_d) + C(year)'
# model3 = sm.OLS.from_formula(formula3, data=data).fit(cov_type='HC3')

# # Combine results into one table
# results_table = summary_col([model1, model2, model3], 
#                             stars=True, 
#                             model_names=["FTA1 Only", "FTA1 & FTA2", "FTA1, FTA2 & FTA3"],
#                             info_dict={'R-squared': lambda x: f"{x.rsquared:.4f}",
#                                        'No. observations': lambda x: f"{int(x.nobs)}"})

# # Save results in LaTeX
# latex_path = '../../output/regression/regression_2.tex'
# with open(latex_path, 'w') as f:
#     f.write(results_table.as_latex())

# # Save results in DOCX
# import docx
# doc = docx.Document()
# doc.add_paragraph("Regression Results for FTA Models")
# doc.add_paragraph(results_table.as_text())
# docx_path = '../../output/regression/regression_2.docx'
# doc.save(docx_path)

# print("Regression results saved in LaTeX and DOCX formats.")

# Ensure 'country_id_o', 'country_id_d'categorical variables
data['country_id_o'] = data['country_id_o'].astype('category')
data['country_id_d'] = data['country_id_d'].astype('category')

# First regression: FTA1 only
formula1 = 'tradeflow_comtrade_o ~ FTA1_ijt + C(country_id_o) + C(country_id_d)'
model1 = sm.OLS.from_formula(formula1, data=data).fit(cov_type='HC3')

# Second regression: FTA1 and FTA2
formula2 = 'tradeflow_comtrade_o ~ FTA1_ijt + FTA2_ict + C(country_id_o) + C(country_id_d)'
model2 = sm.OLS.from_formula(formula2, data=data).fit(cov_type='HC3')

# Third regression: FTA1, FTA2, and FTA3
formula3 = 'tradeflow_comtrade_o ~ FTA1_ijt + FTA2_ict + FTA3_cjt + C(country_id_o) + C(country_id_d)'
model3 = sm.OLS.from_formula(formula3, data=data).fit(cov_type='HC3')

# Combine results into one table
results_table = summary_col([model1, model2, model3], 
                            stars=True, 
                            model_names=["(1)", "(2)", "(3)"],
                            info_dict={'R-squared': lambda x: f"{x.rsquared:.4f}",
                                       'No. observations': lambda x: f"{int(x.nobs)}"})

# Save results in LaTeX
latex_path = '../../output/regression/regression_2.tex'
with open(latex_path, 'w') as f:
    f.write(results_table.as_latex())

# Save results in DOCX
import docx
doc = docx.Document()
doc.add_paragraph("Regression Results for FTA Models")
doc.add_paragraph(results_table.as_text())
docx_path = '../../output/regression/regression_2.docx'
doc.save(docx_path)

print("Regression results saved in LaTeX and DOCX formats.")
